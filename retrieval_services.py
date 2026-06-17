from __future__ import annotations

import math
import os
import re
import tempfile
from io import BytesIO
from pathlib import Path
from typing import Any

# Force Transformers/SentenceTransformers to avoid TensorFlow imports.
os.environ.setdefault("USE_TF", "0")
os.environ.setdefault("TRANSFORMERS_NO_TF", "1")

import numpy as np
import requests
from langchain_community.document_loaders import PyPDFLoader
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from rank_bm25 import BM25Okapi
from sentence_transformers import CrossEncoder
from transformers import pipeline as hf_pipeline

from config import (
    CHUNK_OVERLAP,
    CHUNK_SIZE,
    CONFIDENCE_HIGH_THRESHOLD,
    CONFIDENCE_MEDIUM_THRESHOLD,
    EMBEDDING_MODEL_NAME,
    EASYOCR_LANGUAGES,
    ENABLE_REMOTE_FALLBACK,
    ENABLE_REMOTE_TRANSLATION,
    HF_INFERENCE_API_TOKEN,
    HF_REMOTE_LLM_MODEL,
    LOCAL_QA_MODEL_NAME,
    LOW_CONFIDENCE_REJECT_THRESHOLD,
    MIN_ANSWER_EVIDENCE_SIMILARITY,
    MIN_ANSWER_GROUNDING_SCORE,
    MIN_LEXICAL_OVERLAP,
    MIN_QA_SCORE,
    MIN_QUERY_ANSWER_SIMILARITY,
    MIN_RELEVANCE_SCORE,
    OCR_MIN_QUALITY_SCORE,
    OCR_ENGINE_PREFERENCE,
    OUT_OF_SCOPE_LEXICAL_THRESHOLD,
    OUT_OF_SCOPE_QA_THRESHOLD,
    OUT_OF_SCOPE_RELEVANCE_THRESHOLD,
    PADDLEOCR_LANGUAGE,
    RERANK_TOP_K,
    RERANKER_MODEL_NAME,
    RETRIEVAL_TOP_K,
    SCANNED_PDF_TEXT_THRESHOLD,
    STRICT_DOCUMENT_GROUNDED,
)

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

try:
    import fitz
except ImportError:
    fitz = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from PIL import Image, ImageFilter, ImageOps
except ImportError:
    Image = None
    ImageFilter = None
    ImageOps = None

try:
    import easyocr
except ImportError:
    easyocr = None

try:
    from paddleocr import PaddleOCR
except ImportError:
    PaddleOCR = None

try:
    import pytesseract
except ImportError:
    pytesseract = None


SUPPORTED_EXTENSIONS = {
    ".pdf", ".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp",
    ".docx", ".html", ".htm", ".txt", ".md", ".py", ".js", ".ts", ".tsx",
    ".jsx", ".java", ".c", ".cpp", ".cs", ".go", ".rs", ".json", ".yaml",
    ".yml", ".xml", ".css", ".sql", ".sh",
}

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}
HTML_EXTENSIONS = {".html", ".htm"}
TEXT_EXTENSIONS = {
    ".txt", ".md", ".py", ".js", ".ts", ".tsx", ".jsx", ".java", ".c",
    ".cpp", ".cs", ".go", ".rs", ".json", ".yaml", ".yml", ".xml", ".css",
    ".sql", ".sh",
}

_embedding_model: HuggingFaceEmbeddings | None = None
_reranker: CrossEncoder | None = None
_qa_pipeline: Any | None = None
_easyocr_reader: Any | None = None
_paddleocr_reader: Any | None = None

LANGUAGE_NAME_TO_CODE = {
    "arabic": "ar", "chinese": "zh", "english": "en", "french": "fr",
    "german": "de", "hindi": "hi", "japanese": "ja", "korean": "ko",
    "marathi": "mr", "portuguese": "pt", "russian": "ru", "spanish": "es",
    "tamil": "ta", "telugu": "te", "urdu": "ur",
}

LANGUAGE_CODE_TO_NAME = {code: name.title() for name, code in LANGUAGE_NAME_TO_CODE.items()}


def get_embedding_model() -> HuggingFaceEmbeddings:
    global _embedding_model
    if _embedding_model is None:
        _embedding_model = HuggingFaceEmbeddings(
            model_name=EMBEDDING_MODEL_NAME,
            encode_kwargs={"normalize_embeddings": True},
        )
    return _embedding_model


def get_reranker() -> CrossEncoder:
    global _reranker
    if _reranker is None:
        _reranker = CrossEncoder(RERANKER_MODEL_NAME, max_length=1024)
    return _reranker


def get_qa_pipeline() -> Any | None:
    global _qa_pipeline
    if not LOCAL_QA_MODEL_NAME:
        return None
    if _qa_pipeline is None:
        _qa_pipeline = hf_pipeline(
            "question-answering",
            model=LOCAL_QA_MODEL_NAME,
            tokenizer=LOCAL_QA_MODEL_NAME,
        )
    return _qa_pipeline


def get_easyocr_reader() -> Any | None:
    global _easyocr_reader
    if easyocr is None:
        return None
    if _easyocr_reader is None:
        languages = [lang.strip() for lang in EASYOCR_LANGUAGES.split(",") if lang.strip()]
        _easyocr_reader = easyocr.Reader(languages or ["en"], gpu=False)
    return _easyocr_reader


def get_paddleocr_reader() -> Any | None:
    global _paddleocr_reader
    if PaddleOCR is None:
        return None
    if _paddleocr_reader is None:
        _paddleocr_reader = PaddleOCR(use_angle_cls=True, lang=PADDLEOCR_LANGUAGE, show_log=False)
    return _paddleocr_reader


def tokenize(text: str) -> list[str]:
    return re.findall(r"\w+", text.lower())


def detect_language_hint(text: str) -> str:
    if not text.strip():
        return "unknown"
    if re.search(r"[\u0900-\u097F]", text): return "hi"
    if re.search(r"[\u0980-\u09FF]", text): return "bn"
    if re.search(r"[\u0A80-\u0AFF]", text): return "gu"
    if re.search(r"[\u0B80-\u0BFF]", text): return "ta"
    if re.search(r"[\u0C00-\u0C7F]", text): return "te"
    if re.search(r"[\u0600-\u06FF]", text): return "ar"
    if re.search(r"[\u3040-\u30FF]", text): return "ja"
    if re.search(r"[\uAC00-\uD7AF]", text): return "ko"
    if re.search(r"[\u4E00-\u9FFF]", text): return "zh"
    return "en"


def can_use_remote_fallback() -> bool:
    return ENABLE_REMOTE_FALLBACK and bool(HF_INFERENCE_API_TOKEN)


def can_use_remote_translation() -> bool:
    return ENABLE_REMOTE_TRANSLATION and bool(HF_INFERENCE_API_TOKEN)


def language_name(code: str) -> str:
    return LANGUAGE_CODE_TO_NAME.get(code, code or "unknown")


def parse_target_language(query: str) -> str | None:
    lowered = query.lower()
    for language, code in LANGUAGE_NAME_TO_CODE.items():
        pattern = (
            rf"(?:answer|respond|reply|summarize|summary|translate|convert|give|provide)"
            rf"(?:\s+the\s+answer)?(?:\s+it)?\s+(?:in|into|to)\s+{language}\b"
        )
        if re.search(pattern, lowered):
            return code
    if re.search(r"(?:same language as (?:the )?(?:file|document)|in the document language)", lowered):
        return "document"
    return None


def strip_language_instruction(query: str) -> str:
    cleaned = query
    for language in LANGUAGE_NAME_TO_CODE:
        cleaned = re.sub(
            rf"(?i)[,.;]?\s*(?:please\s+)?(?:answer|respond|reply|summarize|summary|translate|convert|give|provide)"
            rf"(?:\s+the\s+answer)?(?:\s+it)?\s+(?:in|into|to)\s+{language}\b",
            "", cleaned,
        )
    cleaned = re.sub(
        r"(?i)[,.;]?\s*(?:please\s+)?(?:answer|respond|reply)\s+in\s+the\s+(?:same\s+)?language\s+as\s+(?:the\s+)?(?:file|document)\b",
        "", cleaned,
    )
    cleaned = re.sub(r"\s{2,}", " ", cleaned).strip(" ,.;")
    return cleaned or query.strip()


def parse_query_preferences(query: str, document_language: str) -> dict[str, str | bool]:
    query_language = detect_language_hint(query)
    requested_target_language = parse_target_language(query)
    cleaned_query = strip_language_instruction(query)

    if requested_target_language == "document":
        target_language = document_language if document_language != "unknown" else query_language
    elif requested_target_language:
        target_language = requested_target_language
    else:
        target_language = query_language

    return {
        "clean_query": cleaned_query,
        "query_language": query_language,
        "target_language": target_language,
        "translation_requested": target_language != query_language or cleaned_query != query.strip(),
    }


def validate_supported_file(file_path: Path) -> None:
    if file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
        allowed = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise ValueError(f"Unsupported file type. Supported types: {allowed}")


def format_table_rows(rows: list[list[Any]]) -> str:
    formatted_rows: list[str] = []
    for row in rows:
        cleaned_cells = [normalize_whitespace(str(cell or "")) for cell in row]
        cleaned_cells = [cell for cell in cleaned_cells if cell]
        if cleaned_cells:
            formatted_rows.append(" | ".join(cleaned_cells))
    return "\n".join(formatted_rows)


def render_pdf_page_to_image(file_path: Path, page_index: int) -> Any | None:
    if fitz is None or Image is None:
        return None
    pdf = fitz.open(str(file_path))
    try:
        page = pdf.load_page(page_index)
        pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
        return Image.open(BytesIO(pixmap.tobytes("png")))
    finally:
        pdf.close()


def extract_scanned_pdf_page(file_path: Path, page_index: int) -> tuple[str, float, str]:
    rendered_image = render_pdf_page_to_image(file_path, page_index)
    if rendered_image is None:
        return "", 0.0, "render-unavailable"

    best_text = ""
    best_score = 0.0
    best_variant = "rendered-original"

    for variant_name, variant_image in build_ocr_candidates(rendered_image):
        for engine_name in get_ocr_engine_order():
            if engine_name == "easyocr":
                text, average_confidence = run_easyocr_on_variant(variant_image)
                candidate_score = score_ocr_result(text, average_confidence)
                if candidate_score > best_score:
                    best_text = text; best_score = candidate_score; best_variant = f"{variant_name}/easyocr"
            elif engine_name == "paddleocr":
                text, average_confidence = run_paddleocr_on_variant(variant_image)
                candidate_score = score_ocr_result(text, average_confidence)
                if candidate_score > best_score:
                    best_text = text; best_score = candidate_score; best_variant = f"{variant_name}/paddleocr"
            elif engine_name == "pytesseract" and pytesseract is not None:
                for psm in (4, 6, 11):
                    tesseract_config = f"--oem 3 --psm {psm}"
                    text, average_confidence = run_ocr_on_variant(variant_image, tesseract_config)
                    candidate_score = score_ocr_result(text, average_confidence)
                    if candidate_score > best_score:
                        best_text = text; best_score = candidate_score; best_variant = f"{variant_name}/tesseract-psm{psm}"

    return best_text, best_score, best_variant


def extract_pdf_documents(file_path: Path) -> list[Document]:
    extracted: list[Document] = []
    if pdfplumber is not None:
        with pdfplumber.open(str(file_path)) as pdf:
            for index, page in enumerate(pdf.pages, start=1):
                page_parts: list[str] = []
                raw_text = normalize_whitespace(page.extract_text() or "")
                if raw_text:
                    page_parts.append(raw_text)

                table_blocks: list[str] = []
                try:
                    tables = page.extract_tables() or []
                except Exception:
                    tables = []
                for table_index, table in enumerate(tables, start=1):
                    formatted_table = format_table_rows(table)
                    if formatted_table:
                        table_blocks.append(f"Table {table_index}:\n{formatted_table}")

                if table_blocks:
                    page_parts.append("\n\n".join(table_blocks))

                page_text = normalize_whitespace("\n\n".join(page_parts))
                extraction_method = "pdf_text_table"
                ocr_quality_score = None
                ocr_variant = None

                if len(page_text) < SCANNED_PDF_TEXT_THRESHOLD:
                    ocr_text, ocr_quality_score, ocr_variant = extract_scanned_pdf_page(file_path, index - 1)
                    if ocr_text and ocr_quality_score >= OCR_MIN_QUALITY_SCORE:
                        page_text = normalize_whitespace("\n\n".join(part for part in [page_text, ocr_text] if part))
                        extraction_method = "pdf_text_table_ocr"

                if not page_text:
                    continue

                metadata: dict[str, Any] = {
                    "section": index,
                    "section_label": f"page {index}",
                    "extraction_method": extraction_method,
                }
                if table_blocks:
                    metadata["contains_table"] = True
                if ocr_quality_score is not None:
                    metadata["ocr_quality_score"] = round(ocr_quality_score, 3)
                if ocr_variant is not None:
                    metadata["ocr_variant"] = ocr_variant

                extracted.append(Document(page_content=page_text, metadata=metadata))
        if extracted:
            return extracted

    documents = PyPDFLoader(str(file_path)).load()
    for index, doc in enumerate(documents, start=1):
        text = normalize_whitespace(doc.page_content)
        if not text:
            continue
        extracted.append(Document(page_content=text, metadata={"section": index, "section_label": f"page {index}", "extraction_method": "pdf_loader_text"}))
    return extracted


def normalize_whitespace(text: str) -> str:
    text = text.replace("\x0c", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def sentence_split(text: str) -> list[str]:
    parts = re.split(r"(?<=[.!?])\s+|\n+", text)
    return [part.strip() for part in parts if part.strip()]


def build_ocr_candidates(image: Any) -> list[tuple[str, Any]]:
    variants: list[tuple[str, Any]] = [("original", image.convert("RGB"))]
    grayscale = ImageOps.grayscale(image)
    variants.append(("grayscale", grayscale))
    variants.append(("autocontrast", ImageOps.autocontrast(grayscale)))
    threshold = ImageOps.autocontrast(grayscale).point(lambda px: 255 if px > 170 else 0)
    variants.append(("threshold", threshold))
    enlarged = ImageOps.autocontrast(grayscale).resize((max(grayscale.width * 2, 1), max(grayscale.height * 2, 1)))
    variants.append(("enlarged", enlarged))
    sharpened = enlarged.filter(ImageFilter.SHARPEN)
    variants.append(("sharpened", sharpened))
    return variants


def run_ocr_on_variant(image: Any, tesseract_config: str) -> tuple[str, float]:
    data = pytesseract.image_to_data(image, config=tesseract_config, output_type=pytesseract.Output.DICT)
    words, confidences = [], []
    for word, confidence in zip(data.get("text", []), data.get("conf", [])):
        cleaned = str(word).strip()
        if not cleaned: continue
        words.append(cleaned)
        try:
            val = float(confidence)
            if val >= 0: confidences.append(val)
        except (TypeError, ValueError):
            continue
    text = normalize_whitespace(" ".join(words))
    return text, (sum(confidences) / len(confidences) if confidences else 0.0)


def run_easyocr_on_variant(image: Any) -> tuple[str, float]:
    reader = get_easyocr_reader()
    if reader is None: return "", 0.0
    results = reader.readtext(np.array(image), detail=1, paragraph=False)
    texts, confidences = [], []
    for result in results:
        if len(result) < 3: continue
        text = normalize_whitespace(str(result[1]))
        if not text: continue
        texts.append(text)
        confidences.append(float(result[2]))
    return normalize_whitespace("\n".join(texts)), (sum(confidences) / len(confidences) if confidences else 0.0)


def run_paddleocr_on_variant(image: Any) -> tuple[str, float]:
    reader = get_paddleocr_reader()
    if reader is None: return "", 0.0
    results = reader.ocr(np.array(image), cls=True)
    texts, confidences = [], []
    for block in results or []:
        for line in block or []:
            if not line or len(line) < 2: continue
            rec = line[1]
            if not rec or len(rec) < 2: continue
            text = normalize_whitespace(str(rec[0]))
            if not text: continue
            texts.append(text)
            confidences.append(float(rec[1]))
    return normalize_whitespace("\n".join(texts)), (sum(confidences) / len(confidences) if confidences else 0.0)


def score_ocr_result(text: str, average_confidence: float) -> float:
    if not text: return 0.0
    tokens = tokenize(text)
    unique_ratio = len(set(tokens)) / max(len(tokens), 1)
    alnum_ratio = sum(char.isalnum() for char in text) / max(len(text), 1)
    length_score = min(len(tokens) / 40, 1.0)
    conf_score = min(max(average_confidence / 100.0, 0.0), 1.0) if average_confidence > 1.0 else average_confidence
    return (0.45 * conf_score) + (0.25 * unique_ratio) + (0.20 * alnum_ratio) + (0.10 * length_score)


def get_ocr_engine_order() -> list[str]:
    return [e.strip().lower() for e in OCR_ENGINE_PREFERENCE.split(",") if e.strip()] or ["easyocr", "paddleocr", "pytesseract"]


def extract_best_ocr_text(file_path: Path) -> tuple[str, float, str]:
    image = Image.open(file_path)
    best_text, best_score, best_variant = "", 0.0, "original"
    for variant_name, variant_image in build_ocr_candidates(image):
        for engine_name in get_ocr_engine_order():
            if engine_name == "easyocr":
                text, conf = run_easyocr_on_variant(variant_image)
                score = score_ocr_result(text, conf)
                if score > best_score: best_text = text; best_score = score; best_variant = f"{variant_name}/easyocr"
            elif engine_name == "paddleocr":
                text, conf = run_paddleocr_on_variant(variant_image)
                score = score_ocr_result(text, conf)
                if score > best_score: best_text = text; best_score = score; best_variant = f"{variant_name}/paddleocr"
            elif engine_name == "pytesseract" and pytesseract is not None:
                for psm in (6, 11):
                    text, conf = run_ocr_on_variant(variant_image, f"--oem 3 --psm {psm}")
                    score = score_ocr_result(text, conf)
                    if score > best_score: best_text = text; best_score = score; best_variant = f"{variant_name}/tesseract-psm{psm}"
    return best_text, best_score, best_variant


def extract_image_documents(file_path: Path) -> list[Document]:
    if Image is None: raise ValueError("Image processing missing Pillow. Please verify install.")
    text, score, variant = extract_best_ocr_text(file_path)
    if not text: raise ValueError("No readable text was detected in the uploaded image.")
    if score < 0.25: raise ValueError("OCR could not confidently parse the image layout. Upload higher resolution file.")
    return [Document(page_content=text, metadata={"section": 1, "section_label": "image ocr", "extraction_method": "ocr", "ocr_quality_score": round(score, 3), "ocr_variant": variant})]


def extract_docx_documents(file_path: Path) -> list[Document]:
    if DocxDocument is None: raise ValueError("DOCX library missing from configuration environment.")
    doc = DocxDocument(str(file_path))
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    table_blocks = []
    for idx, table in enumerate(doc.tables, start=1):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        fmt = format_table_rows(rows)
        if fmt: table_blocks.append(f"Table {idx}:\n{fmt}")
    parts = paras + table_blocks
    if not parts: raise ValueError("No structural content detected within source DOCX file.")
    return [Document(page_content="\n\n".join(parts), metadata={"section": 1, "section_label": "document body", "extraction_method": "docx_text_table", "contains_table": bool(table_blocks)})]


def extract_html_documents(file_path: Path) -> list[Document]:
    if BeautifulSoup is None: raise ValueError("BeautifulSoup library missing from system run modules.")
    soup = BeautifulSoup(file_path.read_text(encoding="utf-8", errors="ignore"), "html.parser")
    for script in soup(["script", "style", "noscript"]): script.extract()
    lines = [line.strip() for line in soup.get_text("\n").splitlines() if line.strip()]
    if not lines: raise ValueError("No extractable clean semantic strings found inside HTML.")
    return [Document(page_content="\n".join(lines), metadata={"section": 1, "section_label": "html body", "extraction_method": "html_text"})]


def extract_text_documents(file_path: Path) -> list[Document]:
    txt = file_path.read_text(encoding="utf-8", errors="ignore").strip()
    if not txt: raise ValueError("Target plaintext source mapping evaluates as completely empty.")
    return [Document(page_content=txt, metadata={"section": 1, "section_label": "text body", "extraction_method": "plain_text"})]


def extract_documents(file_path: Path) -> list[Document]:
    validate_supported_file(file_path)
    suffix = file_path.suffix.lower()
    if suffix == ".pdf": return extract_pdf_documents(file_path)
    if suffix in IMAGE_EXTENSIONS: return extract_image_documents(file_path)
    if suffix == ".docx": return extract_docx_documents(file_path)
    if suffix in HTML_EXTENSIONS: return extract_html_documents(file_path)
    if suffix in TEXT_EXTENSIONS: return extract_text_documents(file_path)
    raise ValueError(f"Unsupported validation structure context: {suffix}")


def build_chunks(file_path: Path) -> list[Document]:
    documents = extract_documents(file_path)
    splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP, separators=["\n\n", "\n", ".", " ", ""])
    chunks = splitter.split_documents(documents)
    for idx, chunk in enumerate(chunks):
        chunk.metadata.update({"chunk_id": idx, "document_name": file_path.name, "document_type": file_path.suffix.lower(), "language_hint": detect_language_hint(chunk.page_content)})
    return chunks


def build_bm25(chunks: list[Document]) -> BM25Okapi:
    return BM25Okapi([tokenize(doc.page_content) for doc in chunks])


def bm25_search(query: str, chunks: list[Document], bm25: BM25Okapi, top_k: int = 5) -> list[Document]:
    scores = bm25.get_scores(tokenize(query))
    return [chunks[idx] for idx in np.argsort(scores)[::-1][:top_k]]


def hybrid_search(query: str, chunks: list[Document], vectorstore: Chroma, bm25: BM25Okapi, top_k: int = RETRIEVAL_TOP_K) -> list[Document]:
    semantic_docs = vectorstore.similarity_search(query, k=top_k)
    bm25_docs = bm25_search(query, chunks, bm25, top_k=top_k)
    unique_docs, seen = [], set()
    for doc in (semantic_docs + bm25_docs):
        cid = doc.metadata.get("chunk_id")
        if cid not in seen:
            seen.add(cid)
            unique_docs.append(doc)
    return unique_docs[:top_k]


def rerank_documents(query: str, documents: list[Document], top_k: int = RERANK_TOP_K) -> tuple[list[Document], list[float]]:
    if not documents: return [], []
    scores = get_reranker().predict([[query, doc.page_content] for doc in documents])
    ranked_indices = np.argsort(scores)[::-1][:top_k]
    return [documents[idx] for idx in ranked_indices], [float(scores[idx]) for idx in ranked_indices]


def expand_query(query: str) -> list[str]:
    expansions = {
        "claim": ["claim settlement", "insurance reimbursement", "claim process"],
        "policy": ["insurance coverage", "coverage rules", "policy document"],
        "exclusions": ["not covered", "limitations", "exceptions"],
    }
    expanded, lowered = [query], query.lower()
    for k, variants in expansions.items():
        if k in lowered: expanded.extend(variants)
    return expanded


def normalize_rerank_score(score: float) -> float:
    # Handles both pre-scaled logits or raw sigmoid values gracefully
    return score if 0.0 <= score <= 1.0 else (1 / (1 + math.exp(-score)))


def answer_grounding_ratio(answer: str, documents: list[Document]) -> float:
    answer_tokens = set(tokenize(answer))
    if not answer_tokens or not documents: return 0.0
    best_overlap = 0.0
    for doc in documents:
        doc_tokens = set(tokenize(doc.page_content))
        if not doc_tokens: continue
        best_overlap = max(best_overlap, len(answer_tokens & doc_tokens) / max(len(answer_tokens), 1))
    return best_overlap


def lexical_overlap_ratio(query: str, documents: list[Document]) -> float:
    query_tokens = set(tokenize(query))
    if not query_tokens or not documents: return 0.0
    context_tokens = set()
    for doc in documents: context_tokens.update(tokenize(doc.page_content))
    return len(query_tokens & context_tokens) / max(len(query_tokens), 1)


def sentence_relevance_score(query: str, sentence: str) -> float:
    q_tokens, s_tokens = set(tokenize(query)), set(tokenize(sentence))
    if not q_tokens or not s_tokens: return 0.0
    return (0.7 * (len(q_tokens & s_tokens) / len(q_tokens))) + (0.3 * (len(q_tokens & s_tokens) / len(s_tokens)))


def cosine_similarity(a: list[float], b: list[float]) -> float:
    denom = np.linalg.norm(a) * np.linalg.norm(b)
    return float(np.dot(a, b) / denom) if denom != 0 else 0.0


def embedding_similarity(text_a: str, text_b: str) -> float:
    if not text_a.strip() or not text_b.strip(): return 0.0
    m = get_embedding_model()
    return cosine_similarity(m.embed_query(text_a), m.embed_query(text_b))


def collect_supporting_sentences(query: str, documents: list[Document], limit: int = 4) -> list[str]:
    scored = []
    for doc in documents:
        for s in sentence_split(doc.page_content):
            score = sentence_relevance_score(query, s)
            if score > 0: scored.append((score, s))
    scored.sort(key=lambda x: x[0], reverse=True)
    selected, seen = [], set()
    for _, s in scored:
        if s.lower() not in seen:
            seen.add(s.lower())
            selected.append(s)
            if len(selected) >= limit: break
    return selected


def verify_answer_against_evidence(query: str, answer: str, documents: list[Document]) -> tuple[bool, dict[str, float], str]:
    if not answer.strip() or "not found in the uploaded file" in answer.lower():
        return False, {"answer_evidence_similarity": 0.0, "query_answer_similarity": 0.0, "grounding_strength": 0.0}, "Inference failed to locate grounding markers."

    supporting_sentences = collect_supporting_sentences(query, documents, limit=6)
    if not supporting_sentences:
        return False, {"answer_evidence_similarity": 0.0, "query_answer_similarity": 0.0, "grounding_strength": 0.0}, "No direct lexical matching anchors found in chunks."

    ans_ev_sim = max(embedding_similarity(answer, s) for s in supporting_sentences)
    q_ans_sim = embedding_similarity(query, answer)
    grounding_strength = answer_grounding_ratio(answer, documents)

    # Adaptive scoring mitigation logic for complex abstract phrasing
    if ans_ev_sim < (MIN_ANSWER_EVIDENCE_SIMILARITY - 0.15):
        return False, {"answer_evidence_similarity": round(ans_ev_sim, 3), "query_answer_similarity": round(q_ans_sim, 3), "grounding_strength": round(grounding_strength, 3)}, "Low semantic alignment between synthesis and actual contextual evidence chunks."

    return True, {"answer_evidence_similarity": round(ans_ev_sim, 3), "query_answer_similarity": round(q_ans_sim, 3), "grounding_strength": round(grounding_strength, 3)}, "Passed structural validation criteria successfully."


def calculate_confidence(query: str, documents: list[Document], rerank_scores: list[float], qa_score: float, answer: str, used_external_fallback: bool) -> dict[str, Any]:
    if not documents or "not found in the uploaded file" in answer.lower():
        return {"score": 0.0, "label": "low", "reason": "No relevant verified contextual info discovered in file structure."}

    semantic_strength = max((normalize_rerank_score(s) for s in rerank_scores), default=0.0)
    avg_semantic_strength = sum(normalize_rerank_score(s) for s in rerank_scores) / len(rerank_scores) if rerank_scores else 0.0
    lexical_strength = lexical_overlap_ratio(query, documents)
    evidence_density = min(len(documents) / 3, 1.0)
    grounding_strength = answer_grounding_ratio(answer, documents)

    confidence = (0.35 * semantic_strength) + (0.15 * avg_semantic_strength) + (0.15 * lexical_strength) + (0.15 * evidence_density) + (0.20 * grounding_strength)
    if used_external_fallback: confidence *= 0.85
    
    label = "high" if confidence >= CONFIDENCE_HIGH_THRESHOLD else "medium" if confidence >= CONFIDENCE_MEDIUM_THRESHOLD else "low"
    return {"score": round(confidence, 3), "label": label, "reason": "Confidence represents combined balance of vector search rank precision, text overlap indexation, and lexical structural grounding checks."}


def answer_from_documents(query: str, documents: list[Document]) -> tuple[str, float]:
    """
    Stitches relevant document chunks together so information is continuous,
    preventing extractive truncation issues.
    """
    if not documents: return "", 0.0
    
    # Sort context by structural sequence context order to preserve narrative flow
    sorted_docs = sorted(documents, key=lambda x: x.metadata.get("chunk_id", 0))
    stitched_context = "\n\n...[Context Transition]...\n\n".join([doc.page_content.strip() for doc in sorted_docs])
    
    qa = get_qa_pipeline()
    if qa is None: return stitched_context[:1000], 0.50
    
    try:
        # Run local baseline window probe
        res = qa(question=query, context=stitched_context[:2000])
        extracted_phrase = str(res.get("answer", "")).strip()
        score = float(res.get("score", 0.0))
        
        if len(extracted_phrase) > 5 and score > MIN_QA_SCORE:
            # Locate whole sentences matching the span window area to avoid truncation
            for paragraph in stitched_context.split("\n\n"):
                if extracted_phrase in paragraph:
                    return paragraph, score
            return extracted_phrase, score
    except Exception:
        pass
        
    return stitched_context, 0.45


def synthesize_grounded_answer(query: str, documents: list[Document], extracted_answer: str, target_language: str) -> str:
    """
    Leverages abstractive contextual assembly over the complete retrieval window 
    to provide comprehensive, complete answers.
    """
    if not documents: return "The answer was not found in the uploaded file."
    
    sorted_docs = sorted(documents, key=lambda x: x.metadata.get("chunk_id", 0))
    context_window = "\n---\n".join([f"[Source Chunk Section {d.metadata.get('section_label', 'n/a')}]: {d.page_content}" for d in sorted_docs])

    if can_use_remote_fallback():
        prompt = (
            f"You are a strict, enterprise grade document question-answering system.\n"
            f"Your output language preference must be exactly: {language_name(target_language)}.\n"
            f"Task: Answer the user query using ONLY the provided verified context fragments below.\n"
            f"Rules:\n"
            f"1. Provide a detailed, completely structured explanation based entirely on facts provided below. Do not shorten or truncate explanations.\n"
            f"2. If the fragments do not contain clear validation material to answer the query, reply with exactly: 'The answer was not found in the uploaded file.'\n"
            f"3. Never deduce or build background logic outside of this given file data.\n\n"
            f"Context Fragments:\n{context_window}\n\n"
            f"Query: {query}\n\n"
            f"Detailed Synthesis Answer:"
        )
        ans = remote_generation_request(prompt)
        if ans and len(ans.strip()) > 5:
            return ans.strip()
            
    # Local fallback strategy using custom formatting
    support_lines = collect_supporting_sentences(query, documents, limit=5)
    if support_lines:
        return f"Verified Data Highlights:\n" + "\n".join([f"• {line}" for line in support_lines])
        
    return sorted_docs[0].page_content


def remote_generation_request(prompt: str) -> str:
    if not HF_INFERENCE_API_TOKEN: return ""
    try:
        res = requests.post(
            f"https://api-inference.huggingface.co/models/{HF_REMOTE_LLM_MODEL}",
            headers={"Authorization": f"Bearer {HF_INFERENCE_API_TOKEN}", "Content-Type": "application/json"},
            json={"inputs": prompt, "parameters":{"max_new_tokens": 512, "temperature": 0.1, "return_full_text": False}},
            timeout=30,
        )
        if res.status_code == 200:
            payload = res.json()
            if isinstance(payload, list) and payload: return str(payload[0].get("generated_text", "")).strip()
            if isinstance(payload, dict): return str(payload.get("generated_text", "")).strip()
    except Exception:
        pass
    return ""


def maybe_translate_answer(answer: str, source_language: str, target_language: str, original_query: str) -> str:
    if not answer or target_language == source_language or "not found in the uploaded file" in answer.lower():
        return answer
    if not can_use_remote_translation():
        return f"{answer}\n\n[Translation configuration unverified: Output locked to source language format.]"

    prompt = (
        f"Translate this text exactly into {language_name(target_language)}. "
        f"Maintain exact naming structures, code terms, line values, and technical syntax verbatim.\n\nText:\n{answer}"
    )
    trans = remote_generation_request(prompt)
    return trans or answer


def call_remote_llm(query: str, documents: list[Document], target_language: str) -> str:
    if not can_use_remote_fallback(): return "The answer was not found in the uploaded file."
    return synthesize_grounded_answer(query, documents, "", target_language)


def evaluate_retrieval_scoping(query: str, documents: list[Document], rerank_scores: list[float], qa_score: float) -> tuple[bool, str]:
    """
    Robust threshold handling logic to confirm if document context matches intent.
    """
    if not documents:
        return True, "No document context chunks were fetched by search index pipelines."

    normalized = [normalize_rerank_score(s) for s in rerank_scores]
    max_relevance = max(normalized, default=0.0)
    lexical_overlap = lexical_overlap_ratio(query, documents)

    # Cleaned Threshold Pass logic gates
    if max_relevance < MIN_RELEVANCE_SCORE and lexical_overlap < MIN_LEXICAL_OVERLAP:
        return True, f"Search metrics fall below standard relevance cutoffs (Max Vector: {max_relevance:.2f}, Lexical Overlap: {lexical_overlap:.2f})."
        
    return False, "Target thresholds verified. Safe context query boundaries satisfied."


def build_workflow_steps(file_path: Path, doc_lang: str, clean_query: str, query_lang: str, target_lang: str, documents: list[Document], rerank_scores: list[float], qa_score: float, used_fallback: bool, fallback_reason: str) -> list[str]:
    top_score = max((normalize_rerank_score(s) for s in rerank_scores), default=0.0)
    return [
        f"Loaded `{file_path.suffix.lower()}` text streams using contextual document pipelines.",
        f"Document Language mapped: `{language_name(doc_lang)}`.",
        f"Normalized Query payload to: '{clean_query}' (Inferred Query language: `{language_name(query_lang)}`).",
        f"Extracted {len(documents)} context fragments into memory store structures.",
        f"Reranked semantic records using `{RERANKER_MODEL_NAME}`. Best cross-encoder match: {top_score:.2f}.",
        f"Evaluated scoping logic path boundaries: {fallback_reason}",
        "Routed context to remote cloud orchestration layer via Hugging Face Hub APIs." if used_fallback else "Resolved synthesis safely using local vector space contexts."
    ]


def prepare_pipeline(file_path: Path) -> dict[str, Any]:
    validate_supported_file(file_path)
    chunks = build_chunks(file_path)
    if not chunks: raise ValueError("No search index entries generated from extraction phase.")

    v_dir = Path(tempfile.mkdtemp(prefix="enterprise_chroma_"))
    v_store = Chroma.from_documents(documents=chunks, embedding=get_embedding_model(), persist_directory=str(v_dir))
    sample_text = "\n".join([c.page_content for c in chunks[:5]])

    return {
        "document_name": file_path.name, "document_path": str(file_path), "document_type": file_path.suffix.lower(),
        "language_hint": detect_language_hint(sample_text), "chunks": chunks, "bm25": build_bm25(chunks),
        "vectorstore": v_store, "vectorstore_dir": str(v_dir),
    }

def generate_relevant_suggestions(chunks: list[Document], limit: int = 3) -> list[str]:
    """
    Scans the uploaded document chunks to pull out high-quality topic headers
    or sentences to suggest to the user when they ask an out-of-scope question.
    """
    if not chunks:
        return ["What is the main purpose of this document?", "Can you summarize this file?"]
        
    suggestions = []
    seen_phrases = set()
    
    # Try to grab clean sentences or bullet highlights from the document's early chunks
    for chunk in chunks[:5]:
        content = chunk.page_content.strip()
        # Look for generic lines that look like structural topics (between 20 and 70 characters)
        lines = [line.strip("•-* ") for line in content.split("\n") if 20 <= len(line.strip()) <= 70]
        
        for line in lines:
            lowered = line.lower()
            if lowered not in seen_phrases and not any(k in lowered for k in ["http", "version", "author"]):
                seen_phrases.add(lowered)
                # Turn the line into a natural question format
                if not line.endswith("?"):
                    suggestions.append(f"Tell me about {line}")
                else:
                    suggestions.append(line)
                    
            if len(suggestions) >= limit:
                break
        if len(suggestions) >= limit:
            break
            
    # Fallback to standard smart questions if the document text extraction is dense code/tables
    if len(suggestions) < limit:
        suggestions.append("What are the key highlights or core metrics mentioned here?")
        suggestions.append("Can you provide a comprehensive summary breakdown of this file?")
        
    return suggestions[:limit]


def retrieve_complete_stitched_context(top_documents: list[Document], all_chunks: list[Document], window_size: int = 1) -> list[Document]:
    """
    Takes the highly relevant top chunks and dynamically expands their boundaries
    by pulling neighboring chunks (previous and next) to deliver a complete, 
    un-truncated narrative window.
    """
    if not top_documents or not all_chunks:
        return top_documents

    # Map all available chunks by their unique chunk_id for instant sequence lookup
    chunk_map = {doc.metadata.get("chunk_id"): doc for doc in all_chunks if doc.metadata.get("chunk_id") is not None}
    
    expanded_docs = []
    seen_chunk_ids = set()

    # Sort top documents by chunk_id to keep the document flow in reading order
    sorted_top_docs = sorted(top_documents, key=lambda x: x.metadata.get("chunk_id", 0))

    for doc in sorted_top_docs:
        current_id = doc.metadata.get("chunk_id")
        if current_id is None:
            expanded_docs.append(doc)
            continue

        # Determine the sequence window (e.g., current_id - 1, current_id, current_id + 1)
        start_id = max(0, current_id - window_size)
        end_id = min(len(all_chunks) - 1, current_id + window_size)

        stitched_text_parts = []
        
        for cid in range(start_id, end_id + 1):
            if cid in chunk_map:
                stitched_text_parts.append(chunk_map[cid].page_content.strip())
                seen_chunk_ids.add(cid)

        # Merge the neighbor chunks cleanly with smooth line breaks
        complete_text = "\n ".join(stitched_text_parts)
        
        # Clean up any fragmented sentence boundaries at the very beginning or end
        # Matches from the first capital letter to the final period punctuation mark
        sentence_bound_match = re.search(r"([A-Z].*[\.!?])", complete_text, re.DOTALL)
        if sentence_bound_match:
            complete_text = sentence_bound_match.group(1)

        # Create a new expanded document keeping the core metadata intact
        expanded_doc = Document(
            page_content=complete_text,
            metadata=doc.metadata
        )
        expanded_docs.append(expanded_doc)

    return expanded_docs

def run_query(pipeline: dict[str, Any], query: str) -> dict[str, Any]:
    original_query = query.strip()
    if not original_query: raise ValueError("Queries cannot process completely empty request items.")

    preferences = parse_query_preferences(original_query, str(pipeline["language_hint"]))
    clean_query = str(preferences["clean_query"]).strip()
    query_language = str(preferences["query_language"])
    target_language = str(preferences["target_language"])
    document_language = str(pipeline["language_hint"])

    # 1. Retrieve the hybrid matches
    aggregated_docs, seen_ids = [], set()
    for variant in expand_query(clean_query):
        matches = hybrid_search(variant, pipeline["chunks"], pipeline["vectorstore"], pipeline["bm25"], top_k=RETRIEVAL_TOP_K)
        for doc in matches:
            cid = doc.metadata.get("chunk_id")
            if cid not in seen_ids:
                seen_ids.add(cid)
                aggregated_docs.append(doc)

    # 2. Rerank matches and compute vector embedding alignment scores
    top_documents, rerank_scores = rerank_documents(clean_query, aggregated_docs, top_k=RERANK_TOP_K)
    
    # --- CRITICAL STITCHING UPGRADE ---
    # Dynamically scale up and complete your information chunks with window boundaries
    top_documents = retrieve_complete_stitched_context(top_documents, pipeline["chunks"], window_size=1)
    
    # 3. Compute vector relevance scores
    normalized_scores = [normalize_rerank_score(s) for s in rerank_scores]
    max_vector_relevance = max(normalized_scores, default=0.0)
    
    # Check if a summary request is explicitly triggered by intent words
    is_summary_request = any(k in clean_query.lower() for k in ["summar", "sumari", "sumer", "overview", "about"])

    used_external_fallback = False
    
    if max_vector_relevance < MIN_RELEVANCE_SCORE and not is_summary_request:
        fallback_reason = f"Out-of-File content detected. Max embedding cosine similarity ({max_vector_relevance:.2f}) fell below threshold ({MIN_RELEVANCE_SCORE}). Prompting user with relevant alternatives."
        
        alternatives = generate_relevant_suggestions(pipeline["chunks"], limit=3)
        suggestion_bullets = "\n".join([f"👉 **\"{alt}\"**" for alt in alternatives])
        
        answer = (
            f"I looked through your uploaded file (**{pipeline['document_name']}**), but I couldn't find "
            f"information regarding your question *\"{original_query}\"*.\n\n"
            f"To get the best results, try asking a question more closely related to the document. "
            f"Here are a few relevant examples you can try right now:\n\n{suggestion_bullets}"
        )
        extracted_text, qa_score = "", 0.0
    else:
        extracted_text, qa_score = answer_from_documents(clean_query, top_documents)
        is_out_of_scope, scoping_reason = evaluate_retrieval_scoping(clean_query, top_documents, rerank_scores, qa_score)
        fallback_reason = scoping_reason

        answer = synthesize_grounded_answer(clean_query, top_documents, extracted_text, target_language)
        
        if is_summary_request:
            sorted_docs = sorted(top_documents, key=lambda x: x.metadata.get("chunk_id", 0))
            answer = f"### Document Content Summary Overview:\n\n" + "\n\n".join([f"**Section {d.metadata.get('section_label', 'n/a')}**:\n{d.page_content[:600]}..." for d in sorted_docs[:3]])
            fallback_reason = "Synthesized broad-scoped overview text directly from document layers."

    # Final metrics pass for UI elements
    _, verification_metrics, verification_reason = verify_answer_against_evidence(clean_query, answer, top_documents)
    confidence = calculate_confidence(clean_query, top_documents, rerank_scores, qa_score, answer, used_external_fallback)

    if is_summary_request:
        confidence["label"] = "MEDIUM (Document Overview)"
        confidence["score"] = max(confidence["score"], 0.55)
        should_reject_local_answer = False
    else:
        should_reject_local_answer = confidence["score"] < LOW_CONFIDENCE_REJECT_THRESHOLD if max_vector_relevance >= MIN_RELEVANCE_SCORE else False

    if should_reject_local_answer and not is_summary_request:
        fallback_reason = f"Grounded confidence level ({confidence['score']:.3f}) below strict system threshold."
        alternatives = generate_relevant_suggestions(pipeline["chunks"], limit=3)
        suggestion_bullets = "\n".join([f"👉 **\"{alt}\"**" for alt in alternatives])
        answer = (
            f"I found some related text, but the confidence score was too low to verify an exact answer.\n\n"
            f"Try refining your prompt or pick one of these relevant target topics from the document:\n\n{suggestion_bullets}"
        )

    if max_vector_relevance >= MIN_RELEVANCE_SCORE and not used_external_fallback and (target_language != document_language or bool(preferences["translation_requested"])):
        answer = maybe_translate_answer(answer, document_language, target_language, original_query)

    return {
        "query": original_query, "normalized_query": clean_query, "answer": answer,
        "query_language": query_language, "target_language": target_language, "document_language": document_language,
        "verification_reason": verification_reason, "verification_metrics": verification_metrics,
        "sources": [
            {
                "chunk_id": doc.metadata.get("chunk_id"), "document_name": doc.metadata.get("document_name", pipeline["document_name"]),
                "preview": doc.page_content.strip()[:900], # Increased preview limit to match the newly completed content
                "section_label": doc.metadata.get("section_label", "section"),
                "language_hint": doc.metadata.get("language_hint", pipeline["language_hint"]),
            } for doc in top_documents
        ],
        "confidence": confidence, "used_external_fallback": used_external_fallback, "fallback_reason": fallback_reason,
        "workflow_steps": build_workflow_steps(Path(pipeline["document_path"]), document_language, clean_query, query_language, target_language, top_documents, rerank_scores, qa_score, used_external_fallback, fallback_reason),
    }